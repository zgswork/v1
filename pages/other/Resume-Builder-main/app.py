"""简历生成与预览工具 - Flask后端"""
import json
import os
import uuid
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
DATA_DIR = BASE_DIR / "data"
FONTS_DIR = BASE_DIR / "fonts"
OUTPUT_DIR = BASE_DIR / "output"

OUTPUT_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)

# ============ 模板扫描 ============

def scan_templates():
    """扫描 /templates 目录，返回所有模板信息"""
    templates = []
    if not TEMPLATES_DIR.exists():
        return templates
    for folder in sorted(TEMPLATES_DIR.iterdir()):
        if folder.is_dir():
            config_file = folder / "config.json"
            if config_file.exists():
                with open(config_file, "r", encoding="utf-8") as f:
                    config = json.load(f)
                config["id"] = folder.name
                config["path"] = str(folder)
                templates.append(config)
    return templates


def render_resume(template_id: str, data: dict) -> str:
    """使用指定模板渲染简历 HTML"""
    template_dir = TEMPLATES_DIR / template_id
    html_file = template_dir / "index.html"

    if not html_file.exists():
        # Fallback: return error HTML
        return f"<html><body><h1>Template '{template_id}' not found</h1></body></html>"

    with open(html_file, "r", encoding="utf-8") as f:
        html_content = f.read()

    # 简单模板引擎：替换 {{key}} 占位符
    # 支持嵌套：{{profile.name}}, {{workExperiences[0].company}}
    import re

    def resolve(ctx, path):
        """解析 'profile.name' 或 'items[0].name' 路径"""
        keys = re.split(r"\.|\[|\]", path)
        current = ctx
        for key in keys:
            if key == "" or key is None:
                continue
            if key.isdigit():
                if isinstance(current, list) and int(key) < len(current):
                    current = current[int(key)]
                else:
                    return ""
            elif isinstance(current, dict) and key in current:
                current = current[key]
            elif isinstance(current, list):
                return ""
            else:
                return ""
        if isinstance(current, (dict, list)):
            return ""
        return str(current) if current else ""

    def preprocess_data(data):
        """预处理数据：展平嵌套结构供模板使用"""
        processed = {}
        for key, value in data.items():
            if isinstance(value, list) and key not in ("featuredSkills",):
                processed[key] = []
                for item in value:
                    if isinstance(item, dict):
                        new_item = dict(item)
                        descs = item.get("descriptions", [])
                        if isinstance(descs, list) and len(descs) > 0:
                            list_items = []
                            for d in descs:
                                if d and d.strip() and (d.strip()[-1] in (":", "：")) and len(d.strip()) < 30:
                                    list_items.append(f'<li class="desc-label">{d}</li>')
                                else:
                                    list_items.append(f"<li>{d}</li>")
                            new_item["_descriptions_html"] = "".join(list_items)
                            new_item["_list_tag"] = "ul"
                        processed[key].append(new_item)
                    else:
                        processed[key].append(item)
            elif isinstance(value, dict):
                processed[key] = dict(value)
                # 只在 skills 下展平 descriptions
                if key == "skills":
                    if "descriptions" in value and isinstance(value["descriptions"], list):
                        processed["_skills_descriptions"] = value["descriptions"]
                    if "featuredSkills" in value and isinstance(value["featuredSkills"], list):
                        processed["_featuredSkills"] = value["featuredSkills"]
                # 展平 selfEvaluation.descriptions
                if key == "selfEvaluation" and "descriptions" in value:
                    processed["_selfeval_descriptions"] = value["descriptions"]
                # 展平 custom.descriptions
                if key == "custom" and "descriptions" in value:
                    processed["_custom_descriptions"] = value["descriptions"]
            else:
                processed[key] = value
        # 展平 profile 字段供模板 {{#if}} 使用
        if "profile" in data and isinstance(data["profile"], dict):
            for pk, pv in data["profile"].items():
                processed[f"_profile_{pk}"] = pv
        return processed

    # 预处理数据
    data = preprocess_data(data)

    def render(html, context):
        """递归渲染模板，支持循环和条件"""
        # 先处理 {{#each}} 循环
        def resolve_if(ctx, path):
            keys = re.split(r"\.|\[|\]", path)
            current = ctx
            for key in keys:
                if key == "" or key is None:
                    continue
                if key.isdigit():
                    if isinstance(current, list) and int(key) < len(current):
                        current = current[int(key)]
                    else:
                        return ""
                elif isinstance(current, dict) and key in current:
                    current = current[key]
                else:
                    return ""
            return current

        # 处理条件 {{#if key}}...{{/if}} 支持嵌套路径如 profile.email
        def process_loop(match):
            key = match.group(1).strip()
            body = match.group(2)
            items = resolve_if(context, key)
            if not isinstance(items, list):
                items = []
            result = []
            for item in items:
                if isinstance(item, dict):
                    # 合并 context 和 item（item 优先）
                    item_context = {**context, **item}
                elif isinstance(item, str):
                    item_context = {**context, "_value": item}
                else:
                    item_context = context
                # 递归渲染循环体
                result.append(render(body, item_context))
            return "".join(result)

        html = re.sub(
            r"\{\{#each\s+([\w.]+)\}\}(.*?)\{\{/each\}\}",
            process_loop,
            html,
            flags=re.DOTALL,
        )

        # 解析路径用于 {{#if}}，保留 dict/list 类型（与 resolve 不同，不转换类型）
        def process_if(match):
            key = match.group(1).strip()
            body = match.group(2)
            value = resolve_if(context, key)
            if isinstance(value, list):
                return body if len(value) > 0 else ""
            if isinstance(value, dict):
                return body if value else ""
            return body if value else ""

        html = re.sub(
            r"\{\{#if\s+([\w.]+)\}\}(.*?)\{\{/if\}\}",
            process_if,
            html,
            flags=re.DOTALL,
        )

        # 处理简单占位符 {{key}} 和 {{nested.key}} 和 {{nested[0].key}}
        def process_placeholder(match):
            path = match.group(1).strip()
            return resolve(context, path)

        html = re.sub(r"\{\{(.+?)\}\}", process_placeholder, html)

        return html

    html_content = render(html_content, data)

    # 注入CSS
    css_file = template_dir / "style.css"
    css_content = ""
    if css_file.exists():
        with open(css_file, "r", encoding="utf-8") as f:
            css_content = f.read()

    # 最终HTML
    final_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: "Noto Sans SC", "Source Han Sans SC", "Microsoft YaHei", sans-serif; }}
{css_content}
</style>
<style>
/* PDF打印样式 */
@media print {{
  body {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  @page {{ size: A4; margin: 0; }}
}}
</style>
</head>
<body>
{html_content}
</body>
</html>"""

    return final_html


# ============ API路由 ============

@app.route("/")
def index():
    """返回前端编辑器页面"""
    return send_file(BASE_DIR / "static" / "editor.html")


@app.route("/output/<path:filename>")
def serve_output(filename):
    """提供 output 目录下的静态文件"""
    return send_from_directory(OUTPUT_DIR, filename)


@app.route("/data/<path:filename>")
def serve_data(filename):
    """提供 data 目录下的静态文件"""
    return send_from_directory(DATA_DIR, filename)



@app.route("/api/templates")
def api_templates():
    """获取所有模板列表"""
    templates = scan_templates()
    return jsonify(templates)


@app.route("/api/render/<template_id>")
def api_render(template_id):
    """渲染简历预览 HTML"""
    data = request.args.get("data", "{}")
    try:
        resume_data = json.loads(data)
    except json.JSONDecodeError:
        resume_data = {}
    html = render_resume(template_id, resume_data)
    return html


@app.route("/api/preview/<template_id>")
def api_preview(template_id):
    """在iframe中预览简历"""
    data = request.args.get("data", "{}")
    try:
        resume_data = json.loads(data)
    except json.JSONDecodeError:
        resume_data = {}
    html = render_resume(template_id, resume_data)
    return html


@app.route("/api/export-pdf/<template_id>")
def api_export_pdf(template_id):
    """使用 Playwright (Chromium) 导出 PDF，保证与预览完全一致"""
    data = request.args.get("data", "{}")
    try:
        resume_data = json.loads(data)
    except json.JSONDecodeError:
        resume_data = {}

    html = render_resume(template_id, resume_data)

    html_path = OUTPUT_DIR / f"temp_{uuid.uuid4().hex}.html"
    pdf_path = OUTPUT_DIR / f"resume_{uuid.uuid4().hex}.pdf"

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)

    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(f"file:///{html_path.absolute().as_posix()}")
            page.wait_for_load_state("networkidle")
            page.pdf(
                path=str(pdf_path),
                format="A4",
                print_background=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"}
            )
            browser.close()
    except Exception as e:
        html_path.unlink(missing_ok=True)
        # Fallback: open the HTML in browser with print dialog
        print_html = html.replace("</body>", "<script>window.onload=function(){window.print();}</script></body>")
        return print_html

    html_path.unlink(missing_ok=True)
    return send_file(pdf_path, as_attachment=True, download_name=f"resume.pdf",
                     mimetype="application/pdf")


@app.route("/api/sample-data")
def api_sample_data():
    """获取示例简历数据"""
    sample_file = DATA_DIR / "sample.json"
    if sample_file.exists():
        with open(sample_file, "r", encoding="utf-8") as f:
            return jsonify(json.load(f))
    # 返回默认示例数据
    sample = {
        "profile": {
            "name": "张三",
            "email": "zhangsan@example.com",
            "phone": "13800138000",
            "location": "广州市天河区",
            "title": "高级前端工程师",
            "summary": "8年Web前端开发经验，精通React/Vue/TypeScript，主导过多个大型项目..."
        },
        "workExperiences": [
            {
                "company": "某科技有限公司",
                "jobTitle": "高级前端工程师",
                "date": "2021.03 - 至今",
                "descriptions": [
                    "负责公司核心产品前端架构设计与开发",
                    "带领5人前端团队，推动组件化与工程化建设",
                    "优化首屏加载性能，LCP从4.2s降至1.8s"
                ]
            },
            {
                "company": "某互联网公司",
                "jobTitle": "前端开发工程师",
                "date": "2018.07 - 2021.02",
                "descriptions": [
                    "参与电商平台前端开发，使用React+TypeScript",
                    "开发可复用组件库，提升团队开发效率30%"
                ]
            }
        ],
        "educations": [
            {
                "school": "中山大学",
                "degree": "本科 · 计算机科学与技术",
                "date": "2014.09 - 2018.06",
                "descriptions": ["GPA 3.8/4.0，获校级奖学金"]
            }
        ],
        "skills": {
            "featuredSkills": [
                {"skill": "React/Vue", "rating": 5},
                {"skill": "TypeScript", "rating": 5},
                {"skill": "Node.js", "rating": 4},
                {"skill": "CSS/HTML", "rating": 5}
            ],
            "descriptions": ["Webpack/Vite", "Git/GitHub", "Docker", "CI/CD"]
        },
        "projects": [
            {
                "project": "企业级中台管理系统",
                "date": "2022.03 - 2022.12",
                "descriptions": [
                    "基于React + Ant Design Pro搭建",
                    "实现权限管理、数据可视化、多语言等模块",
                    "支撑日均10万+用户访问"
                ]
            }
        ],
        "custom": {
            "descriptions": ["个人博客(blog.example.com)累计输出技术文章50+篇", "GitHub开源项目star 2k+"]
        }
    }
    return jsonify(sample)


@app.route("/api/export-word/<template_id>")
def api_export_word(template_id):
    """使用 python-docx 导出 Word 文档"""
    data = request.args.get("data", "{}")
    try:
        resume_data = json.loads(data)
    except json.JSONDecodeError:
        resume_data = {}

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # Profile
    profile = resume_data.get("profile", {})
    if profile.get("name"):
        h = doc.add_heading(profile["name"], level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if profile.get("title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile["title"])
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Contact info
    contacts = [v for k, v in profile.items() if v and k in ("email","phone","location")]
    if contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contacts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

    if profile.get("summary"):
        doc.add_heading("个人简介", level=1)
        doc.add_paragraph(profile["summary"])

    # Work Experience
    works = resume_data.get("workExperiences", [])
    if works:
        doc.add_heading("工作经历", level=1)
        for work in works:
            title_text = f"{work.get('company','')} | {work.get('jobTitle','')} | {work.get('date','')}"
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.bold = True
            for desc in work.get("descriptions", []):
                doc.add_paragraph(desc, style='List Bullet')

    # Education
    edus = resume_data.get("educations", [])
    if edus:
        doc.add_heading("教育经历", level=1)
        for edu in edus:
            title_text = f"{edu.get('school','')} | {edu.get('degree','')} | {edu.get('date','')}"
            doc.add_paragraph(title_text, style='List Bullet')

    # Projects
    projects = resume_data.get("projects", [])
    if projects:
        doc.add_heading("项目经历", level=1)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(f"{proj.get('project','')} | {proj.get('date','')}")
            run.bold = True
            for desc in proj.get("descriptions", []):
                doc.add_paragraph(desc, style='List Bullet')

    # Skills
    skills = resume_data.get("skills", {})
    if skills.get("featuredSkills") or skills.get("descriptions"):
        doc.add_heading("技能", level=1)
        fskills = skills.get("featuredSkills", [])
        if fskills:
            skill_names = [s.get("skill","") for s in fskills if s.get("skill")]
            doc.add_paragraph(" | ".join(skill_names))
        descs = skills.get("descriptions", [])
        if descs:
            doc.add_paragraph(", ".join(descs))

    # Custom
    custom = resume_data.get("custom", {}).get("descriptions", [])
    if custom:
        doc.add_heading("其他信息", level=1)
        for c in custom:
            doc.add_paragraph(c, style='List Bullet')

    # Save
    docx_path = OUTPUT_DIR / f"resume_{uuid.uuid4().hex}.docx"
    doc.save(str(docx_path))
    return send_file(docx_path, as_attachment=True, download_name=f"resume.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    print("Resume Builder starting...")
    print("Template dir:", TEMPLATES_DIR)
    print("Visit: http://localhost:5000")
    app.run(debug=True, port=5000)
